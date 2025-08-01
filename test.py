import os
import gradio as gr
import asyncio
from lightrag import LightRAG, QueryParam
from lightrag.llm.ollama import ollama_model_complete, ollama_embed
from lightrag.kg.shared_storage import initialize_pipeline_status
from lightrag.utils import EmbeddingFunc, setup_logger
import re
from docx import Document

setup_logger("lightrag", level="INFO")

WORKING_DIR = "./chat_rag_storage"
os.makedirs(WORKING_DIR, exist_ok=True)
print("\U0001F4C1 Danh sách file sau khi xóa:", os.listdir(WORKING_DIR))

def fetch_uploaded_gradio_file(file) -> str:
    try:
        if file.name.endswith(".docx"):
            doc = Document(file.name)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            return '\n'.join(full_text)
        with open(file.name, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        raise Exception(f"Lỗi đọc file: {e}")

def convert_docx_to_txt(file):
    try:
        doc = Document(file.name)
        full_text = [para.text for para in doc.paragraphs if para.text.strip()]
        txt_content = '\n'.join(full_text)
        output_path = file.name.replace(".docx", ".txt")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(txt_content)
        return output_path
    except Exception as e:
        return f"❌ Error: {e}"

def split_into_sentences(text):
    sentence_endings = re.compile(r'(?<=[.!?])\s+')
    return [s.strip() for s in sentence_endings.split(text) if s.strip()]

def split_paragraph(paragraph, max_len=8192):
    sentences = split_into_sentences(paragraph)
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 1 > max_len:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
        current_chunk += sentence + " "

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return chunks

def answer_score(answer, content):
    try:
        from evaluate import load
        bertscore = load("bertscore")
        result = bertscore.compute(predictions=[answer], references=[content], lang="en")
        return f"Score: {result['f1'][0]:.4f}"
    except Exception as e:
        return f"❌ Error in BERTScore: {str(e)}"

async def initialize_chatbot_rag():
    rag = LightRAG(
        working_dir=WORKING_DIR,
        llm_model_func=ollama_model_complete,
        llm_model_name="llama3.1:8b",
        embedding_func=EmbeddingFunc(
            embedding_dim=768,
            max_token_size=8192,
            func=lambda texts: ollama_embed(texts, embed_model="nomic-embed-text"),
        )
    )
    await rag.initialize_storages()
    await initialize_pipeline_status()
    return rag

async def add_document_to_rag(file, rag: LightRAG):
    try:
        content = fetch_uploaded_gradio_file(file)
        success_count = 0
        failed_paragraphs = []

        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        paragraphs = list(set(paragraphs))
        for i, paragraph in enumerate(paragraphs, 1):
            if paragraph.strip() == "":
                continue
            try:
                if len(paragraph) > 8192:
                    sub_chunks = split_paragraph(paragraph)
                    for sub in sub_chunks:
                        sub = "Only provide information strictly extracted from the source documents. Do not add personal opinions, interpretations, or any additional sentences beyond the original content." + sub.strip()
                        await rag.ainsert(sub)
                        success_count += 1
                        print(f"✅ Đã thêm đoạn con {success_count}")
                else:
                    await rag.ainsert("Only provide information strictly extracted from the source documents. Do not add personal opinions, interpretations, or any additional sentences beyond the original content. If several parts of the document convey the same meaning, only include the first matching passage." + paragraph)
                    success_count += 1
                    print(f"✅ Đã thêm đoạn {success_count}")
            except Exception as line_error:
                failed_paragraphs.append(f"Đoạn {i}: {line_error}")
                print(f"❌ Lỗi đoạn {i}")

        if failed_paragraphs:
            msg = f"⚠️ Thêm {success_count} đoạn thành công. {len(failed_paragraphs)} đoạn lỗi:\n" + "\n".join(failed_paragraphs[:5])
            return msg, False
        else:
            return f"✅ Đã thêm {success_count} đoạn tài liệu thành công!", True

    except Exception as e:
        return f"❌ Lỗi khi đọc file: {e}", False

async def RAG_chatbot(message, history, rag, file):
    history = history or []
    content = fetch_uploaded_gradio_file(file) if file else ""
    prompt = f"""
    Based on the internal document, extract information related to the topic below:
    - DO NOT ADD ANY PERSONAL OPINION.
    - Only return content that is directly found in the document.
    - Only return examples that are explicitly stated in the internal document. Do not invent new examples or scenarios, even if they might be helpful.
    - If the document does not contain any information related to the topic, return "No information found in the document.
    - DO NOT USE ADDITIONAL INFOMATION LIKE "References","Created date" or "Entity name" OR EVEN "from the Knowledge Graph" in the response.
    - If multiple references have similar meaning, only return the first one.
    Topic: {message}
    """.strip()

    response = await rag.aquery(prompt, QueryParam(mode="local", top_k=10))
    score_result = answer_score(response, content)

    history.append({"role": "user", "content": message})
    history.append({"role": "assistant", "content": response + f"\n\n📊 {score_result} (if got input file else it will be 0)"})

    return history, history

def create_rag_wrapper(rag):
    async def rag_wrapper(message, history, file):
        return await RAG_chatbot(message, history, rag, file)
    return rag_wrapper

async def main():
    rag = await initialize_chatbot_rag()

    with gr.Blocks() as demo:
        gr.Markdown("## 🧠 Chatbot tài liệu dùng LightRAG + Ollama")
        rag_state = gr.State(rag)
        is_rag_ready = gr.State(False)
        chatbot_ui = gr.Chatbot(type="messages")
        msg = gr.Textbox(label="💬 Nhập câu hỏi")
        history = gr.State([])

        with gr.Row():
            send_btn = gr.Button("📩 Gửi")
            upload_btn = gr.Button("📥 Thêm tài liệu")

        file_input = gr.File(label="📄 Upload file .txt/.md/.docx", file_types=[".txt", ".md", ".docx"])
        upload_result = gr.Textbox(label="Kết quả thêm tài liệu", interactive=False)

        upload_btn.click(fn=add_document_to_rag, inputs=[file_input, rag_state], outputs=[upload_result, is_rag_ready])

        rag_wrapper_func = create_rag_wrapper(rag)
        send_btn.click(fn=rag_wrapper_func, inputs=[msg, history, file_input], outputs=[chatbot_ui, history])
        msg.submit(fn=rag_wrapper_func, inputs=[msg, history, file_input], outputs=[chatbot_ui, history])

        demo.launch(server_name="127.0.0.1", server_port=9621, show_api=False, share=True)

if __name__ == "__main__":
    import nest_asyncio
    nest_asyncio.apply()
    asyncio.run(main())